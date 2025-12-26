"""
Generate Auckland Rain Monitoring System Documentation
"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
from datetime import datetime

def create_document():
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    title = doc.add_heading('Auckland Council', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Rain Monitoring System Documentation', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('INTERNAL DOCUMENT').bold = True
    
    doc.add_paragraph()
    
    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta2.add_run(f'Generated: {datetime.now().strftime("%d %B %Y")}')
    
    doc.add_paragraph()
    meta3 = doc.add_paragraph()
    meta3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta3.add_run('Version 1.0 - DRAFT')
    
    doc.add_page_break()
    
    # =========================================================================
    # TABLE OF CONTENTS (placeholder)
    # =========================================================================
    doc.add_heading('Table of Contents', level=1)
    toc = doc.add_paragraph()
    toc.add_run('1. Introduction\n')
    toc.add_run('2. System Architecture\n')
    toc.add_run('3. Rain Gauge Monitoring\n')
    toc.add_run('4. Rain Radar Monitoring\n')
    toc.add_run('5. ARI Calculation Methodology\n')
    toc.add_run('6. Alarm System\n')
    toc.add_run('7. Technical Reference\n')
    toc.add_run('8. Appendices\n')
    
    doc.add_page_break()
    
    # =========================================================================
    # 1. INTRODUCTION
    # =========================================================================
    doc.add_heading('1. Introduction', level=1)
    
    doc.add_heading('1.1 Purpose', level=2)
    doc.add_paragraph(
        'This document provides comprehensive documentation of Auckland Council\'s '
        'rain monitoring system, which consists of two complementary data sources: '
        'rain gauges providing point measurements and rain radar providing spatial '
        'coverage across stormwater catchments.'
    )
    
    doc.add_heading('1.2 Scope', level=2)
    doc.add_paragraph(
        'This documentation covers:'
    )
    bullets = doc.add_paragraph()
    bullets.add_run('• Rain gauge network configuration and alarm setup\n')
    bullets.add_run('• Rain radar (QPE) data collection and processing\n')
    bullets.add_run('• ARI (Annual Recurrence Interval) calculation methodology\n')
    bullets.add_run('• Alarm trigger logic and historical analysis\n')
    bullets.add_run('• Technical reference for data access via Moata API')
    
    doc.add_heading('1.3 Definitions and Acronyms', level=2)
    
    # Definitions table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Term'
    hdr_cells[1].text = 'Definition'
    
    definitions = [
        ('ARI', 'Annual Recurrence Interval - the average time between events of a given magnitude'),
        ('QPE', 'Quantitative Precipitation Estimate - radar-derived rainfall estimates'),
        ('TP108', 'Technical Publication 108 - NIWA methodology for rainfall frequency analysis'),
        ('Moata', 'Auckland Council\'s asset management and monitoring platform'),
        ('Trace', 'A time series data stream from a sensor or calculated value'),
        ('Threshold', 'A configured value that triggers an alarm when exceeded'),
        ('Catchment', 'A geographic area that drains to a common outlet'),
    ]
    
    for term, defn in definitions:
        row = table.add_row().cells
        row[0].text = term
        row[1].text = defn
    
    doc.add_page_break()
    
    # =========================================================================
    # 2. SYSTEM ARCHITECTURE
    # =========================================================================
    doc.add_heading('2. System Architecture', level=1)
    
    doc.add_heading('2.1 Overview', level=2)
    doc.add_paragraph(
        'The rain monitoring system integrates data from multiple sources through '
        'the Moata platform, which provides unified access to sensor data, alarm '
        'configuration, and historical records.'
    )
    
    doc.add_heading('2.2 Data Flow', level=2)
    doc.add_paragraph(
        'Data flows through the system as follows:'
    )
    flow = doc.add_paragraph()
    flow.add_run('1. Data Collection: Rain gauges and radar systems collect precipitation data\n')
    flow.add_run('2. Data Ingestion: Raw data is ingested into the Moata platform\n')
    flow.add_run('3. Processing: Data is processed, including ARI calculations for radar data\n')
    flow.add_run('4. Alarm Evaluation: Processed values are compared against configured thresholds\n')
    flow.add_run('5. Notification: Exceedances trigger alarms and notifications')
    
    doc.add_heading('2.3 Data Sources Summary', level=2)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Source'
    hdr[1].text = 'Coverage'
    hdr[2].text = 'Resolution'
    hdr[3].text = 'Primary Use'
    
    sources = [
        ('Rain Gauges', '76 active stations', '5-minute intervals', 'Point measurements, ARI alarms'),
        ('Rain Radar', '233 catchments\n25,298 pixels', '1-minute intervals', 'Spatial coverage, area exceedance'),
    ]
    
    for source in sources:
        row = table.add_row().cells
        for i, val in enumerate(source):
            row[i].text = val
    
    doc.add_page_break()
    
    # =========================================================================
    # 3. RAIN GAUGE MONITORING
    # =========================================================================
    doc.add_heading('3. Rain Gauge Monitoring', level=1)
    
    doc.add_heading('3.1 Network Coverage', level=2)
    doc.add_paragraph(
        'The rain gauge network consists of physical sensors deployed across the '
        'Auckland region. Not all gauges in the system are actively monitored for '
        'alarms.'
    )
    
    # Coverage table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Category'
    hdr[1].text = 'Count'
    
    coverage = [
        ('Total gauges in dataset', '264'),
        ('Excluded: Non-Auckland (by keyword)', '67'),
        ('Excluded: No physical sensor (forecast/nowcast only)', '44'),
        ('Excluded: No recent telemetered data', '77'),
        ('Active Auckland rain gauges', '76'),
    ]
    
    for cat, count in coverage:
        row = table.add_row().cells
        row[0].text = cat
        row[1].text = count
    
    doc.add_heading('3.2 Filtering Criteria', level=2)
    doc.add_paragraph(
        'Gauges are filtered based on the following criteria to identify active '
        'Auckland stations:'
    )
    
    criteria = doc.add_paragraph()
    criteria.add_run('Step 1 - Auckland Location: ').bold = True
    criteria.add_run('Gauge name must contain Auckland-related keywords (ACC, Auckland, '
                    'NSCC, Metrowater, etc.) and must not contain exclusion keywords '
                    '(WSL, Watercare, WDC, etc.)\n\n')
    criteria.add_run('Step 2 - Physical Sensor: ').bold = True
    criteria.add_run('Gauge must have a physical rainfall sensor trace (not just '
                    'forecast or nowcast data)\n\n')
    criteria.add_run('Step 3 - Recent Data: ').bold = True
    criteria.add_run('Gauge must have telemetered data within the last 3 months')
    
    doc.add_heading('3.3 Trace Types', level=2)
    doc.add_paragraph(
        'Each rain gauge can have multiple data traces. The following trace types '
        'are present in the network:'
    )
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Trace Type'
    hdr[1].text = 'Description'
    
    traces = [
        ('Rainfall (primary)', 'Raw rainfall measurement from sensor'),
        ('Rainfall 60 min window sum', 'Rolling 60-minute rainfall total'),
        ('Max TP108 ARI', 'Maximum ARI value across all durations'),
        ('Rain Nowcast', 'Short-term rainfall prediction'),
        ('Rain Forecast', 'Longer-term rainfall forecast'),
    ]
    
    for trace, desc in traces:
        row = table.add_row().cells
        row[0].text = trace
        row[1].text = desc
    
    doc.add_heading('3.4 Alarm Configuration', level=2)
    doc.add_paragraph(
        'Alarms are configured on the rain gauge network to detect significant '
        'rainfall events and data quality issues.'
    )
    
    # Alarm config table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Alarm Type'
    hdr[1].text = 'Count'
    hdr[2].text = 'Description'
    
    alarms = [
        ('Overflow Monitoring', '694', 'Threshold-based alarms on rainfall/ARI values'),
        ('Data Recency', '76', 'Alerts when data stops flowing from a gauge'),
    ]
    
    for alarm in alarms:
        row = table.add_row().cells
        for i, val in enumerate(alarm):
            row[i].text = val
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('ARI Threshold: ').bold = True
    p.add_run('The standard threshold for ARI-based alarms is ')
    p.add_run('5 years').bold = True
    p.add_run('. This means an alarm is triggered when the calculated ARI exceeds '
             '5 years, indicating rainfall intensity that would be expected to occur '
             'on average once every 5 years.')
    
    doc.add_page_break()
    
    # =========================================================================
    # 4. RAIN RADAR MONITORING
    # =========================================================================
    doc.add_heading('4. Rain Radar Monitoring', level=1)
    
    doc.add_heading('4.1 Spatial Coverage', level=2)
    doc.add_paragraph(
        'Rain radar provides spatial coverage of precipitation across the Auckland '
        'region. Data is organized by stormwater catchments, with each catchment '
        'mapped to multiple radar pixels.'
    )
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'Value'
    
    metrics = [
        ('Stormwater catchments', '233'),
        ('Total radar pixels', '25,298'),
        ('Average pixels per catchment', '~109'),
        ('Pixels with TP108 coefficients', '19,356'),
    ]
    
    for metric, val in metrics:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = val
    
    doc.add_heading('4.2 Data Characteristics', level=2)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Parameter'
    hdr[1].text = 'Value'
    
    params = [
        ('Data Type', 'QPE (Quantitative Precipitation Estimate)'),
        ('Temporal Resolution', '1 minute'),
        ('Spatial Resolution', 'Radar pixel (~1km²)'),
        ('TraceSet Collection ID', '1'),
        ('TraceSet ID', '3'),
        ('API Limit - Pixels per request', '150 (recommended: 50)'),
        ('API Limit - Time range', '24 hours per request'),
    ]
    
    for param, val in params:
        row = table.add_row().cells
        row[0].text = param
        row[1].text = val
    
    doc.add_heading('4.3 Pixel-to-Catchment Mapping', level=2)
    doc.add_paragraph(
        'Each stormwater catchment is associated with a set of radar pixels that '
        'intersect its geographic boundary. This mapping is:'
    )
    bullets = doc.add_paragraph()
    bullets.add_run('• Static - pixel indices do not change over time\n')
    bullets.add_run('• Determined by geometry intersection via the Moata API\n')
    bullets.add_run('• Cached locally to avoid repeated API calls\n')
    bullets.add_run('• Stored in JSON and pickle formats for efficient reuse')
    
    doc.add_page_break()
    
    # =========================================================================
    # 5. ARI CALCULATION METHODOLOGY
    # =========================================================================
    doc.add_heading('5. ARI Calculation Methodology', level=1)
    
    doc.add_heading('5.1 TP108 Background', level=2)
    doc.add_paragraph(
        'The TP108 methodology, developed by NIWA (National Institute of Water and '
        'Atmospheric Research), provides a standardized approach for calculating '
        'Annual Recurrence Intervals (ARI) from rainfall data. The methodology '
        'accounts for regional variations in rainfall patterns through pixel-specific '
        'coefficients.'
    )
    
    doc.add_heading('5.2 Formula', level=2)
    p = doc.add_paragraph()
    p.add_run('The ARI is calculated using the following formula:\n\n')
    p.add_run('ARI = exp(m × D + b)').bold = True
    p.add_run('\n\nWhere:\n')
    p.add_run('• ARI = Annual Recurrence Interval (years)\n')
    p.add_run('• D = Rainfall depth (mm) accumulated over the duration\n')
    p.add_run('• m = Slope coefficient (specific to pixel and duration)\n')
    p.add_run('• b = Intercept coefficient (specific to pixel and duration)')
    
    doc.add_heading('5.3 Durations', level=2)
    doc.add_paragraph(
        'ARI is calculated for eight standard durations. Shorter durations capture '
        'intense bursts while longer durations capture sustained rainfall events.'
    )
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Duration'
    hdr[1].text = 'Minutes'
    hdr[2].text = 'Typical Depth for ARI=5*'
    
    durations = [
        ('10 minutes', '10', '~13 mm'),
        ('20 minutes', '20', '~21 mm'),
        ('30 minutes', '30', '~26 mm'),
        ('1 hour', '60', '~36 mm'),
        ('2 hours', '120', '~49 mm'),
        ('6 hours', '360', '~75 mm'),
        ('12 hours', '720', '~97 mm'),
        ('24 hours', '1440', '~119 mm'),
    ]
    
    for dur in durations:
        row = table.add_row().cells
        for i, val in enumerate(dur):
            row[i].text = val
    
    doc.add_paragraph()
    doc.add_paragraph(
        '*Note: Depths shown are approximate and vary by pixel location. Values shown '
        'are for a representative Auckland pixel.'
    ).italic = True
    
    doc.add_heading('5.4 Catchment-Level Alarming', level=2)
    doc.add_paragraph(
        'For stormwater catchments, the alarm system monitors the proportion of the '
        'catchment area where ARI exceeds the threshold. This is calculated by:'
    )
    steps = doc.add_paragraph()
    steps.add_run('1. Computing ARI for each pixel within the catchment\n')
    steps.add_run('2. Identifying pixels where ARI exceeds the threshold (5 years)\n')
    steps.add_run('3. Calculating the proportion of total catchment pixels exceeding\n')
    steps.add_run('4. Triggering alarm if proportion exceeds configured threshold')
    
    doc.add_page_break()
    
    # =========================================================================
    # 6. ALARM SYSTEM
    # =========================================================================
    doc.add_heading('6. Alarm System', level=1)
    
    doc.add_heading('6.1 Alarm Types', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Overflow Monitoring\n').bold = True
    p.add_run('Triggered when a measured or calculated value exceeds a configured '
             'threshold. For rain gauges, this typically monitors the "Max TP108 ARI" '
             'trace with a threshold of 5 years.\n\n')
    p.add_run('Data Recency\n').bold = True
    p.add_run('Triggered when data from a sensor has not been received within a '
             'configured time period. This helps identify sensor failures or '
             'communication issues.')
    
    doc.add_heading('6.2 Historical Alarm Analysis', level=2)
    doc.add_paragraph(
        'Analysis of historical alarm events from May to December 2025 shows the '
        'following patterns:'
    )
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Month'
    hdr[1].text = 'Alarm Events'
    
    months = [
        ('May 2025', '5'),
        ('June 2025', '8'),
        ('July 2025', '1'),
        ('August 2025', '2'),
        ('September 2025', '1'),
        ('November 2025', '11'),
        ('December 2025', '10'),
        ('Total', '38'),
    ]
    
    for month, count in months:
        row = table.add_row().cells
        row[0].text = month
        row[1].text = count
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Note: Higher alarm counts in November and December 2025 correspond to '
        'increased rainfall activity during the spring/summer storm season.'
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # 7. TECHNICAL REFERENCE
    # =========================================================================
    doc.add_heading('7. Technical Reference', level=1)
    
    doc.add_heading('7.1 Moata API Endpoints', level=2)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Endpoint'
    hdr[1].text = 'Purpose'
    
    endpoints = [
        ('GET /v1/projects/{projectId}/assets', 'Get assets (gauges, catchments) with optional geometry'),
        ('GET /v1/assets/traces', 'Get traces for one or more assets'),
        ('GET /v1/traces/{traceId}/data/utc', 'Get time series data for a trace'),
        ('GET /v1/traces/{traceId}/thresholds', 'Get alarm thresholds for a trace'),
        ('GET /v1/alarms/overflow-detailed-info-by-trace', 'Get alarm configuration for a trace'),
        ('GET /v1/trace-set-collections/{id}/trace-sets/data', 'Get radar data for pixels'),
        ('GET /v1/trace-set-collections/{id}/pixel-mappings/intersects-geometry', 'Get pixels for a geometry'),
    ]
    
    for ep, purpose in endpoints:
        row = table.add_row().cells
        row[0].text = ep
        row[1].text = purpose
    
    doc.add_heading('7.2 Key Configuration Values', level=2)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Parameter'
    hdr[1].text = 'Value'
    
    config = [
        ('Project ID', '594'),
        ('Rain Gauge Asset Type ID', '25'),
        ('Stormwater Catchment Asset Type ID', '3541'),
        ('TraceSet Collection ID (Radar)', '1'),
        ('TraceSet ID (QPE)', '3'),
        ('ARI Threshold', '5 years'),
        ('Coordinate System (SRID)', '4326 (WGS84)'),
    ]
    
    for param, val in config:
        row = table.add_row().cells
        row[0].text = param
        row[1].text = val
    
    doc.add_heading('7.3 Data Pipeline Scripts', level=2)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Script'
    hdr[1].text = 'Purpose'
    
    scripts = [
        ('retrieve_rain_gauges.py', 'Collect rain gauge data, traces, and alarm configs'),
        ('retrieve_rain_radar.py', 'Collect radar data for stormwater catchments'),
        ('analyze_rain_gauges.py', 'Analyze and filter rain gauge data, generate reports'),
    ]
    
    for script, purpose in scripts:
        row = table.add_row().cells
        row[0].text = script
        row[1].text = purpose
    
    doc.add_page_break()
    
    # =========================================================================
    # 8. APPENDICES
    # =========================================================================
    doc.add_heading('8. Appendices', level=1)
    
    doc.add_heading('Appendix A: Output File Structure', level=2)
    
    structure = doc.add_paragraph()
    structure.add_run('outputs/\n').bold = True
    structure.add_run('├── rain_gauges/\n')
    structure.add_run('│   ├── raw/\n')
    structure.add_run('│   │   └── rain_gauges_traces_alarms.json\n')
    structure.add_run('│   └── analyze/\n')
    structure.add_run('│       ├── alarm_summary.csv\n')
    structure.add_run('│       ├── alarm_summary_full.csv\n')
    structure.add_run('│       ├── all_traces.csv\n')
    structure.add_run('│       ├── analysis_report.txt\n')
    structure.add_run('│       └── dashboard.html\n')
    structure.add_run('└── rain_radar/\n')
    structure.add_run('    └── raw/\n')
    structure.add_run('        ├── catchments/\n')
    structure.add_run('        │   └── stormwater_catchments.csv\n')
    structure.add_run('        ├── pixel_mappings/\n')
    structure.add_run('        │   ├── catchment_pixel_mapping.json\n')
    structure.add_run('        │   └── catchment_pixel_mapping.pkl\n')
    structure.add_run('        ├── radar_data/\n')
    structure.add_run('        │   └── {catchment_id}_{name}.csv\n')
    structure.add_run('        └── collection_summary.json')
    
    doc.add_heading('Appendix B: Glossary of Trace Descriptions', level=2)
    doc.add_paragraph(
        'The following trace descriptions are found in the rain gauge network. '
        'See the generated alarm_summary_full.csv for the complete list with '
        'associated alarm configurations.'
    )
    
    doc.add_heading('Appendix C: Catchment Inventory', level=2)
    doc.add_paragraph(
        'A complete list of 233 stormwater catchments with pixel counts is '
        'available in the collection_summary.json output file.'
    )
    
    # =========================================================================
    # SAVE DOCUMENT
    # =========================================================================
    output_path = Path("outputs/documentation/Rain_Monitoring_System_Documentation.docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    
    print(f"Document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_document()
